import pandas as pd


from pathlib import Path
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P

from langchain_core.documents import Document



class DocxTableExtractor:
    """
    
    """

    def __init__(self):
        """
        
        """
        pass

    def _is_toc_paragraph(self, paragraph: DocxParagraph) -> bool:
        """
        Checking if a paragraph is part of the Table of Contents (TOC).
        """

        if not paragraph.text.strip():
            return False
        if paragraph.style and 'TOC' in paragraph.style.name.upper():
            return True
        return False

    def iter_block_items(self, doc):
        for child in doc.element.body.iterchildren(): # I go through all XML childs of the document body
            if isinstance(child, CT_P): # If it's a paragraph
                yield DocxParagraph(child, doc)
            elif isinstance(child, CT_Tbl): # If it's a table (XML tag)
                yield DocxTable(child, doc)


    def _sanitize_cell_text(self, s: str) -> str:
        if not s:
            return ""
        s = s.replace("\n", " ").replace("\r", " ")
        s = s.replace("|", r"\|")
        return " ".join(s.split())

    def _table_to_markdown(self, tbl: DocxTable) -> str:
        rows = []
        for r in tbl.rows:
            rows.append([self._sanitize_cell_text(c.text) for c in r.cells])

        if not rows:
            return ""

        header = None
        if rows and any(rows[0]):
            header = rows[0]
            rows = rows[1:]

        df = pd.DataFrame(rows, columns=header if header else None)

        try:
            return df.to_markdown(index=False)
        except Exception:
            return df.to_csv(index=False)
        
    def _load_docx_with_table_exclusion(self, file_path: Path) -> tuple[list[Document], list[Document]]:
        """
        """

        doc = DocxDocument(str(file_path))
        blocks = list(self.iter_block_items(doc))


        text_docs: list[Document] = []
        table_docs: list[Document] = []
        buffer: list[str] = []

        def flush_text():
            # We take first the buffered text lines and create a Document.
            if not buffer:
                return
            
            txt = "\n".join([t for t in buffer if t.strip()]).strip()
            buffer.clear()
            if txt:
                text_docs.append(Document(
                    page_content=txt,
                    metadata={
                        "file_name": file_path.name,
                        "file_path": str(file_path),
                        "content_type": "text",
                    }
                ))

        for i, blk in enumerate (blocks):
            if isinstance(blk, DocxParagraph): # It's a paragraph
                # First I check if it's a TOC paragraph
                if self._is_toc_paragraph(blk):
                    continue  # Skip TOC paragraphs

                line = (blk.text or "").strip()
                if line:
                    buffer.append(line) # We accumulate text lines

            elif isinstance(blk, DocxTable): # It's a table
                flush_text()
                md = self._table_to_markdown(blk)
                if md.strip():
                    table_docs.append(Document(
                        page_content=md,
                        metadata={
                            "file_name": file_path.name,  # Au lieu de str(file_path)
                            "file_path": str(file_path),
                            "content_type": "table", 
                        }
                    ))

        flush_text()
        return text_docs, table_docs

def extract_tables_from_docx(
    docx_path: Path,
) -> tuple[list[Document], list[Document]]:
    """
    Extract tables from a DOCX file and return both text and table documents.
    """
    extractor = DocxTableExtractor()
    return extractor._load_docx_with_table_exclusion(docx_path)

if __name__ == "__main__":
    # Test avec un fichier DOCX contenant des tables
    # Remplacez par le chemin de votre fichier DOCX
    docx_path = Path(r"E:\tableaux_test.docx")
    
    if docx_path.exists():
        print(f"Test d'extraction de tables depuis: {docx_path}")
        print("=" * 60)
        
        extractor = DocxTableExtractor()
        text_docs, table_docs = extractor._load_docx_with_table_exclusion(docx_path)
        
        print(f"\nDocuments texte extraits: {len(text_docs)}")
        print("-" * 60)
        for i, doc in enumerate(text_docs):
            print(f"\nTexte {i+1}:")
            print(f"  Fichier: {doc.metadata['file_name']}")
            print(f"  Type: {doc.metadata['content_type']}")
            print(f"\nContenu (extrait):\n{doc.page_content}...\n")
        
        print(f"\nTables extraites: {len(table_docs)}")
        print("-" * 60)
        for i, doc in enumerate(table_docs):
            print(f"\nTable {i+1}:")
            print(f"  Fichier: {doc.metadata['file_name']}")
            print(f"  Type: {doc.metadata['content_type']}")
            print(f"\nContenu:\n{doc.page_content}\n")
        
        print("=" * 60)
        print(f"Total: {len(text_docs)} document(s) texte, {len(table_docs)} table(s)")
    else:
        print(f"ERREUR: Le fichier {docx_path} n'existe pas!")
        print("\nUtilisation:")
        print("  1. Placez un fichier DOCX avec des tables dans le dossier")
        print("  2. Modifiez la variable 'docx_path' avec le bon chemin")
        print("  3. Lancez: python -m app.pipeline.docx_table_extractor")